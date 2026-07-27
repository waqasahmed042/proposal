"""
app.py — Concept Engineers Fee Proposal + Custom E-Signing
(DocuSeal fully removed — self-managed signing via SQLite + reportlab/pypdf)
"""

import os
import tempfile
import shutil
from datetime import date, datetime, timedelta
from pathlib import Path

# Load .env BEFORE importing modules that read env vars at import time
# (auth.py and mailer.py both read MS_* / SECRET_KEY on import).
try:
    from dotenv import load_dotenv

    load_dotenv()
    print("✅ .env file loaded successfully!")
except ImportError:
    print("⚠️  python-dotenv not installed. Run: pip install python-dotenv")

from flask import Flask, request, send_file, render_template, jsonify, abort
from docxtpl import DocxTemplate
from docx import Document as DocxDocument

import db
import signing
import mailer
import auth
import sys

HERE = Path(__file__).parent
TEMPLATE = HERE / "master_proposal_template.docx"

# Where generated + signed PDFs live (local disk for now)
PROPOSALS_DIR = HERE / "storage" / "proposals"
SIGNED_DIR = HERE / "storage" / "signed"
PROPOSALS_DIR.mkdir(parents=True, exist_ok=True)
SIGNED_DIR.mkdir(parents=True, exist_ok=True)

app = Flask(__name__)

# Base URL used to build the signing link sent to clients AND the OAuth
# redirect URI. e.g. https://proposal.conceptengineers.com.au
PUBLIC_BASE_URL = os.getenv("PUBLIC_BASE_URL", "http://localhost:5000")

# Session cookie signing key — REQUIRED for staff sign-in to work.
#
# If unset, we fall back to a random per-process key so the app still boots
# for local dev. But in production that is a TRAP: gunicorn runs multiple
# workers, each would generate a DIFFERENT random key, so a session cookie
# signed by one worker can't be read by another — sign-in fails with
# "Sign-in session expired" because the OAuth flow stored at /auth/start
# isn't found at /auth/callback (a different worker handles it).
# → Always set a fixed SECRET_KEY in .env in production.
_secret_key = os.getenv("SECRET_KEY")
if not _secret_key:
    _secret_key = os.urandom(32).hex()
    print(
        "⚠️  SECRET_KEY is not set in .env — using a random per-process key. "
        "Staff sign-in will be UNRELIABLE with multiple workers "
        "(intermittent 'Sign-in session expired'). Set SECRET_KEY in .env."
    )
app.config["SECRET_KEY"] = _secret_key
app.config["PUBLIC_BASE_URL"] = PUBLIC_BASE_URL

# Harden the session cookie. SESSION_COOKIE_SECURE is on unless we're clearly
# running locally over http (PUBLIC_BASE_URL starts with http://localhost).
_is_local = PUBLIC_BASE_URL.startswith(
    "http://localhost"
) or PUBLIC_BASE_URL.startswith("http://127.0.0.1")
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=not _is_local,
    PERMANENT_SESSION_LIFETIME=timedelta(hours=12),
)

app.register_blueprint(auth.auth_bp)


# Make the signed-in staff user available to every template as `user`.
@app.context_processor
def inject_user():
    return {"user": auth.current_user()}


if not auth.AUTH_ENABLED:
    print("⚠️  AUTH_ENABLED=false — staff pages are UNPROTECTED (dev mode only).")
elif not auth.is_configured():
    print(
        "⚠️  Staff sign-in NOT configured (missing MS_TENANT_ID / MS_CLIENT_ID / "
        "MS_CLIENT_SECRET). Staff pages will be blocked until it is set."
    )
else:
    print("🔒 Staff sign-in enabled (Microsoft Entra ID).")


db.init_db()


def inject_deferred_signing_tags(docx_path: str) -> None:
    """
    Converts [[signature]] / [[date]] (inert to Jinja — docxtpl never
    touches them since they aren't valid Jinja syntax) into literal
    "{{signature}}" / "{{date}}" text inside the already-rendered .docx.

    Must run AFTER tpl.render()+tpl.save(), BEFORE docx_to_pdf(), so these
    tokens survive PDF conversion untouched for signing.py to find and
    replace later, at actual signing time (when the value is known).

    client_name / project_type / project_address / fee are NOT handled
    here — those are known at proposal-generation time and are filled
    directly via normal Jinja context (see build_context() /
    smart_quote_sign()'s ctx dict). Only signature and date are deferred,
    since they're only known once the client actually signs.
    """
    TAG_MAP = {"[[signature]]": "{{signature}}", "[[date]]": "{{date}}"}

    doc = DocxDocument(docx_path)

    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            full_text = "".join(r.text for r in p.runs)
            changed = False
            for src, dst in TAG_MAP.items():
                if src in full_text:
                    full_text = full_text.replace(src, dst)
                    changed = True
            if changed:
                for r in p.runs:
                    r.text = ""
                if p.runs:
                    p.runs[0].text = full_text
                else:
                    p.add_run(full_text)

    replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    doc.save(docx_path)


@app.route("/api/send-quote-invite", methods=["POST"])
@auth.login_required
def send_quote_invite():
    """
    Internal API: create a tokenized Smart Quote link and email it to the client.
    Body: { client_name, client_email, project_address?, note?, project_type? }
    """
    data = request.json or {}
    client_name = (data.get("client_name") or "").strip()
    client_email = (data.get("client_email") or "").strip().lower()
    project_address = (data.get("project_address") or "").strip() or None
    note = (data.get("note") or "").strip() or None
    project_type = (data.get("project_type") or "").strip() or None

    if not client_name:
        return jsonify({"error": "client_name is required"}), 400
    if not client_email or "@" not in client_email:
        return jsonify({"error": "A valid client_email is required"}), 400

    invite = db.create_quote_invite(
        client_name=client_name,
        client_email=client_email,
        project_address=project_address,
        note=note,
        project_type=project_type,
    )
    quote_url = f"{PUBLIC_BASE_URL.rstrip('/')}/quote/{invite['token']}"

    mail_ok, mail_err = mailer.send_quote_invite_email(
        client_name=client_name,
        client_email=client_email,
        quote_url=quote_url,
        project_address=project_address,
        note=note,
    )
    if not mail_ok:
        print(f"[send-quote-invite] ⚠️ email failed: {mail_err}")

    print(f"[send-quote-invite] ✅ client={client_email} url={quote_url}")

    return jsonify(
        {
            "success": True,
            "quoteUrl": quote_url,
            "token": invite["token"],
            "expiresAt": invite["expires_at"],
            "emailSent": mail_ok,
            "emailError": mail_err if not mail_ok else None,
        }
    )


@app.route("/quote/<token>")
def quote_builder(token):
    """Public client-facing quote builder — reached via emailed invite link."""
    row = db.get_quote_invite(token)
    valid, reason = db.is_invite_valid(row)

    if row is None:
        abort(404)

    return render_template(
        "smart_quote.html",
        token=token,
        valid=valid,
        reason=reason,
        client_name=row["client_name"],
        client_email=row["client_email"],
        project_address=row["project_address"] or "",
        project_type=row["project_type"] or "",
    )


@app.route("/api/resend-quote/<token>", methods=["POST"])
@auth.login_required
def resend_quote_invite(token):
    row = db.get_quote_invite(token)
    if row is None:
        return jsonify({"error": "Quote invite not found"}), 404
    if row["status"] == "completed":
        return jsonify({"error": "This quote has already been completed"}), 400
    if row["status"] == "void":
        return jsonify({"error": "This quote link has been cancelled"}), 400

    quote_url = f"{PUBLIC_BASE_URL.rstrip('/')}/quote/{token}"
    mail_ok, mail_err = mailer.send_quote_invite_email(
        client_name=row["client_name"],
        client_email=row["client_email"],
        quote_url=quote_url,
        project_address=row["project_address"],
        note=row["note"],
    )
    if mail_ok:
        return jsonify(
            {"success": True, "message": f"Quote link resent to {row['client_email']}"}
        )
    return jsonify({"error": mail_err}), 500


# Helpers─

# Reference fees from the client's real 26000-FP01-Proposal.docx fee tables.
# Used as defaults when the staff form doesn't override a figure.
REFERENCE_FEES = {
    "da_fee": 9500.0,
    "dd_fee": 14000.0,
    "dd_construction_fee": 3000.0,
    "dd_as_constructed_fee": 1000.0,
    "uu_minor_fee": 2000.0,
    "uu_minor_construction_fee": 1500.0,
    "uu_minor_as_constructed_fee": 1500.0,
    "uu_major_fee": 3000.0,
    "uu_major_as_constructed_fee": 2000.0,
    "unitywater_fee": 4000.0,
}


def _money(n) -> str:
    return f"${n:,.2f}"


def _format_proposal_date(raw: str | None) -> str:
    """
    Converts the ISO date (YYYY-MM-DD) submitted by an <input type="date">
    into the document's display format (e.g. "16 July 2026"). Falls back to
    today's date if the field is empty or not a valid ISO date.
    """
    if raw:
        try:
            return datetime.strptime(raw, "%Y-%m-%d").strftime("%d %B %Y")
        except ValueError:
            return raw  # already a display string (e.g. posted by a test/API caller)
    return date.today().strftime("%d %B %Y")


def build_template_context(
    *,
    contact_name,
    company_name,
    project_address,
    contact_email,
    project_type,
    proposal_ref,
    revision,
    proposal_date,
    tier_label,
    is_da,
    is_dd,
    is_uu_minor,
    is_uu_major,
    is_unitywater,
    has_construction,
    has_as_constructed,
    fees,
    sender_name,
    sender_title,
    sender_email,
    sender_phone,
    contact_phone="",
    accounts_phone="",
    accounts_email="",
) -> dict:
    """
    Builds the docxtpl context for master_proposal_template.docx
    (built from the client's real 26000-FP01-Proposal.docx).

    `fees` is a numeric dict using REFERENCE_FEES keys; totals for each fee
    table are computed here, honouring the construction / as-constructed
    phase toggles.
    """
    f = {**REFERENCE_FEES, **{k: float(v) for k, v in fees.items() if v is not None}}

    dd_total = (
        f["dd_fee"]
        + (f["dd_construction_fee"] if has_construction else 0)
        + (f["dd_as_constructed_fee"] if has_as_constructed else 0)
    )
    uu_minor_total = (
        f["uu_minor_fee"]
        + (f["uu_minor_construction_fee"] if has_construction else 0)
        + (f["uu_minor_as_constructed_fee"] if has_as_constructed else 0)
    )
    uu_major_total = f["uu_major_fee"] + (
        f["uu_major_as_constructed_fee"] if has_as_constructed else 0
    )

    if is_da and is_dd:
        da_dd_label = "Both"
    elif is_da:
        da_dd_label = "DA"
    elif is_dd:
        da_dd_label = "Post-DA"
    else:
        da_dd_label = "—"

    if is_uu_minor or is_uu_major:
        uu_uw_label = "UU"
    elif is_unitywater:
        uu_uw_label = "UW"
    else:
        uu_uw_label = "NONE"

    grand_total = (
        (f["da_fee"] if is_da else 0)
        + (dd_total if is_dd else 0)
        + (uu_minor_total if is_uu_minor else 0)
        + (uu_major_total if is_uu_major else 0)
        + (f["unitywater_fee"] if is_unitywater else 0)
    )

    return {
        # Client / project
        "contact_name": contact_name,
        "company_name": company_name,
        "project_address": project_address,
        "contact_email": contact_email,
        "contact_phone": contact_phone,
        "accounts_phone": accounts_phone,
        "accounts_email": accounts_email,
        "project_type": project_type,
        # Document meta
        "proposal_ref": proposal_ref,
        "revision": revision,
        "proposal_date": proposal_date,
        # Offering
        "tier_label": tier_label,
        "da_dd_label": da_dd_label,
        "uu_uw_label": uu_uw_label,
        # Scope switches
        "is_da": is_da,
        "is_dd": is_dd,
        "is_uu_minor": is_uu_minor,
        "is_uu_major": is_uu_major,
        "is_unitywater": is_unitywater,
        "has_construction": has_construction,
        "has_as_constructed": has_as_constructed,
        # Fees (formatted)
        "da_fee": _money(f["da_fee"]),
        "dd_fee": _money(f["dd_fee"]),
        "dd_construction_fee": _money(f["dd_construction_fee"]),
        "dd_as_constructed_fee": _money(f["dd_as_constructed_fee"]),
        "dd_total": _money(dd_total),
        "uu_minor_fee": _money(f["uu_minor_fee"]),
        "uu_minor_construction_fee": _money(f["uu_minor_construction_fee"]),
        "uu_minor_as_constructed_fee": _money(f["uu_minor_as_constructed_fee"]),
        "uu_minor_total": _money(uu_minor_total),
        "uu_major_fee": _money(f["uu_major_fee"]),
        "uu_major_as_constructed_fee": _money(f["uu_major_as_constructed_fee"]),
        "uu_major_total": _money(uu_major_total),
        "unitywater_fee": _money(f["unitywater_fee"]),
        "unitywater_total": _money(f["unitywater_fee"]),
        # Overall figure (not shown in a table, but stored/reused elsewhere)
        "fee": f"{grand_total:,.2f}",
        "fee_total": grand_total,
        # Sender
        "sender_name": sender_name,
        "sender_title": sender_title,
        "sender_email": sender_email,
        "sender_phone": sender_phone,
    }


def build_context(form) -> dict:
    """Template context from the staff proposal form (/new-proposal).

    Accepts both the new per-scope fields and the legacy checkbox names
    (is_uu → UU Minor, is_uw → UnityWater) so older forms degrade gracefully.
    """

    def money(s, default=None):
        try:
            return float(str(s).replace(",", "").replace("$", ""))
        except Exception:
            return default

    is_da = "is_da" in form
    is_dd = "is_dd" in form
    is_uu_minor = "is_uu_minor" in form or (
        "is_uu" in form and form.get("uu_level", "Minor") != "Major"
    )
    is_uu_major = "is_uu_major" in form or (
        "is_uu" in form and form.get("uu_level") == "Major"
    )
    is_unitywater = "is_unitywater" in form or "is_uw" in form

    fees = {key: money(form.get(key), REFERENCE_FEES[key]) for key in REFERENCE_FEES}

    return build_template_context(
        contact_name=form.get("contact_name", ""),
        company_name=form.get("company_name", ""),
        project_address=form.get("project_address", ""),
        contact_email=form.get("contact_email", ""),
        contact_phone=form.get("contact_phone", ""),
        accounts_phone=form.get("accounts_phone", ""),
        accounts_email=form.get("accounts_email", ""),
        project_type=form.get("project_type", "") or "Residential Subdivision",
        proposal_ref=form.get("proposal_ref", ""),
        revision=form.get("revision", "Rev A"),
        proposal_date=_format_proposal_date(form.get("proposal_date")),
        tier_label=form.get("tier_label", "Gold"),
        is_da=is_da,
        is_dd=is_dd,
        is_uu_minor=is_uu_minor,
        is_uu_major=is_uu_major,
        is_unitywater=is_unitywater,
        has_construction="has_construction" in form,
        has_as_constructed="has_as_constructed" in form,
        fees=fees,
        sender_name=form.get("sender_name", "Concept Engineers Team"),
        sender_title=form.get("sender_title", ""),
        sender_email=form.get("sender_email", "admin@conceptengineers.com.au "),
        sender_phone=form.get("sender_phone", "+61 7 3505 6498"),
    )


def docx_to_pdf(docx_path: str, pdf_path: str) -> tuple[bool, str]:
    """
    Convert DOCX to PDF.

    Windows/macOS: docx2pdf (MS Word) in a fresh subprocess — avoids COM
    threading issues inside Flask's request-handling thread.

    Linux (VPS): docx2pdf requires Word, so go straight to LibreOffice
    headless. Each run gets its own user-profile dir so concurrent
    conversions don't fight over LibreOffice's profile lock.
    """
    import subprocess

    docx2pdf_error = "skipped (requires MS Word — not available on this platform)"

    if sys.platform in ("win32", "darwin"):
        script = (
            "from docx2pdf import convert; " f"convert(r'{docx_path}', r'{pdf_path}')"
        )
        try:
            result = subprocess.run(
                [sys.executable, "-c", script],
                capture_output=True,
                text=True,
                timeout=120,
            )
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                return True, ""
            docx2pdf_error = (
                f"produced no output. stdout={result.stdout.strip()} "
                f"stderr={result.stderr.strip()[:300]}"
            )
        except Exception as e:
            docx2pdf_error = f"{type(e).__name__}: {e}"
        print(f"[docx_to_pdf] docx2pdf failed, trying LibreOffice: {docx2pdf_error}")

    # LibreOffice headless (primary on Linux, fallback elsewhere)
    try:
        out_dir = os.path.dirname(os.path.abspath(pdf_path)) or "."
        with tempfile.TemporaryDirectory(prefix="lo_profile_") as profile_dir:
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    f"-env:UserInstallation={Path(profile_dir).as_uri()}",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    out_dir,
                    docx_path,
                ],
                capture_output=True,
                text=True,
                timeout=120,
            )
        # soffice writes <outdir>/<docx basename>.pdf
        produced = os.path.join(
            out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        )
        if os.path.exists(produced) and os.path.getsize(produced) > 0:
            if os.path.abspath(produced) != os.path.abspath(pdf_path):
                shutil.move(produced, pdf_path)
            return True, ""
        return False, (
            f"docx2pdf: {docx2pdf_error} | "
            f"LibreOffice: {result.stderr.strip() or result.stdout.strip()}"
        )
    except FileNotFoundError:
        return False, (
            f"docx2pdf: {docx2pdf_error} | LibreOffice (soffice) not found on PATH — "
            "install it with: apt install -y --no-install-recommends libreoffice-writer"
        )
    except Exception as e2:
        return False, f"docx2pdf: {docx2pdf_error} | LibreOffice error: {e2}"


# Pages
@app.route("/")
@auth.login_required
def home():
    stats = db.get_stats()
    year = date.today().year
    return render_template("home.html", stats=stats, year=year)


def _default_proposal_ref() -> str:
    """Starting point for the Proposal Ref field, e.g. 'CE-2026-0000'.
    Staff replace the 0000 with the real job number."""
    return f"CE-{date.today().year}-0000"


@app.route("/new-proposal")
@auth.login_required
def index():
    # ISO format (YYYY-MM-DD) is required for <input type="date"> to
    # recognise the value and show it as selected — anything else is
    # silently ignored by the browser, leaving the picker blank.
    today = date.today().isoformat()
    return render_template(
        "form.html", today=today, default_ref=_default_proposal_ref()
    )


@app.route("/smart-quote")
@auth.login_required
def smart_quote_invite_page():
    """Internal page: the team sends a self-service quote link to a client."""
    return render_template("smart_quote_invite.html")


@app.route("/smart-quote-preview")
@auth.login_required
def quote_funnel_preview():
    """
    Internal preview: view the client-facing quote funnel UI directly,
    without needing a real invite token. Uses placeholder client data —
    nothing submitted here is saved (invite_token is not a real invite,
    so /api/smart-quote-sign will reject a submission from this page).
    """
    return render_template(
        "smart_quote.html",
        valid=True,
        token="preview",
        client_name="",
        client_email="",
        project_address="",
        project_type="",
    )


@app.route("/new-proposal-dynamic")
@auth.login_required
def smart_quote():
    """Internal page: the team sends a self-service quote link to a client."""
    return render_template(
        "new_proposal_dynamic.html",
        today=date.today().isoformat(),
        default_ref=_default_proposal_ref(),
    )


# Document generation
def _render_proposal_to(docx_path: Path, ctx: dict) -> None:
    """Render the master template with `ctx` to `docx_path` and inject the
    deferred [[signature]]/[[date]] signing tags. Shared by the download and
    generate-and-send flows."""
    tpl = DocxTemplate(str(TEMPLATE))
    tpl.render(ctx)
    tpl.save(str(docx_path))
    inject_deferred_signing_tags(str(docx_path))


@app.route("/generate", methods=["POST"])
@auth.login_required
def generate():
    """Direct download of a filled proposal (PDF or DOCX). Used for previewing
    a proposal without sending it — the signing flow uses /api/generate-and-send."""
    ctx = build_context(request.form)
    ref = ctx["proposal_ref"] or "proposal"

    tmp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_docx.close()
    _render_proposal_to(Path(tmp_docx.name), ctx)

    if request.form.get("output_fmt", "pdf") == "docx":
        return send_file(
            tmp_docx.name,
            as_attachment=True,
            download_name=f"{ref}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    tmp_pdf = tmp_docx.name.replace(".docx", ".pdf")
    converted, convert_err = docx_to_pdf(tmp_docx.name, tmp_pdf)
    if converted:
        os.unlink(tmp_docx.name)
        return send_file(
            tmp_pdf,
            as_attachment=True,
            download_name=f"{ref}.pdf",
            mimetype="application/pdf",
        )

    print(f"[/generate] PDF conversion failed, falling back to DOCX: {convert_err}")
    return send_file(
        tmp_docx.name,
        as_attachment=True,
        download_name=f"{ref}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# Custom e-signing API
@app.route("/api/generate-and-send", methods=["POST"])
@auth.login_required
def generate_and_send():
    """
    One-request "Generate & Send for Signing" flow.

    Takes the proposal form (multipart — small, just field values), renders the
    proposal, converts it to PDF, stores it, creates a signing token, and emails
    the client the signing link — all server-side.

    This replaces the old two-step /generate?base64 → /api/create-signing-link
    flow, which shipped the whole rendered document (tens of MB with the real
    branded template) to the browser and straight back — slow, and large enough
    to trip nginx's request-size limit (413). Here the document never leaves the
    server, so the request body stays tiny regardless of template size.
    """
    form = request.form
    proposal_ref = (form.get("proposal_ref") or "proposal").strip()
    client_name = (form.get("contact_name") or "").strip()
    client_email = (form.get("contact_email") or "").strip().lower()
    project_address = (form.get("project_address") or "").strip()
    sender_name = (form.get("sender_name") or "Concept Engineers Team").strip()
    sender_title = (form.get("sender_title") or "").strip()
    sender_email = (form.get("sender_email") or "").strip().lower()

    if not client_name:
        return jsonify({"error": "Contact name is required"}), 400
    if not client_email:
        return jsonify({"error": "Contact email is required"}), 400
    if not project_address:
        return jsonify({"error": "Project address is required"}), 400

    safe_ref = (
        "".join(c for c in proposal_ref if c.isalnum() or c in "-_") or "proposal"
    )
    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    docx_path = PROPOSALS_DIR / f"{safe_ref}_{timestamp}.docx"
    pdf_path = PROPOSALS_DIR / f"{safe_ref}_{timestamp}.pdf"

    # Render straight into storage — no base64, no browser round-trip.
    _render_proposal_to(docx_path, build_context(form))

    converted, convert_err = docx_to_pdf(str(docx_path), str(pdf_path))
    if not converted:
        print(f"[generate-and-send] PDF conversion failed: {convert_err}")
        return (
            jsonify({"error": f"Could not convert proposal to PDF: {convert_err}"}),
            500,
        )

    token_info = db.create_signing_token(
        proposal_ref=proposal_ref,
        client_name=client_name,
        client_email=client_email,
        project_address=project_address,
        pdf_path=str(pdf_path),
        sender_name=sender_name,
        sender_title=sender_title,
        sender_email=sender_email,
    )
    signing_url = f"{PUBLIC_BASE_URL.rstrip('/')}/sign/{token_info['token']}"

    print(
        f"[generate-and-send] ✅ ref={proposal_ref} client={client_email} url={signing_url}"
    )

    mail_ok, mail_err = mailer.send_signing_link_email(
        client_name=client_name,
        client_email=client_email,
        proposal_ref=proposal_ref,
        project_address=project_address,
        signing_url=signing_url,
        sender_name=sender_name,
        sender_title=sender_title,
    )
    if not mail_ok:
        print(f"[generate-and-send] ⚠️ email to client failed: {mail_err}")

    return jsonify(
        {
            "success": True,
            "signingUrl": signing_url,
            "token": token_info["token"],
            "expiresAt": token_info["expires_at"],
            "emailSent": mail_ok,
            "emailError": mail_err if not mail_ok else None,
            "message": f"Signing link created for {client_email}",
        }
    )


@app.route("/sign/<token>", methods=["GET"])
def signing_page(token):
    """Public signing page — no login required."""
    row = db.get_token_row(token)
    valid, reason = db.is_token_valid(row)

    if row is None:
        abort(404)

    return render_template(
        "sign.html",
        token=token,
        valid=valid,
        reason=reason,
        proposal_ref=row["proposal_ref"],
        client_name=row["client_name"],
        project_address=row["project_address"],
        status=row["status"],
    )


@app.route("/sign/<token>/document", methods=["GET"])
def signing_document(token):
    """Serves the (unsigned) PDF for preview on the signing page."""
    row = db.get_token_row(token)
    if row is None:
        abort(404)
    pdf_path = row["signed_pdf_path"] or row["pdf_path"]
    if not os.path.exists(pdf_path):
        abort(404)
    return send_file(pdf_path, mimetype="application/pdf")


@app.route("/document/<token>", methods=["GET"])
def document_download(token):
    """
    View or download the proposal document.

    Query params:
      fmt=pdf (default) | docx   — PDF is the signed version once signed;
                                   DOCX is always the original (unsigned) Word file.
      download=1                 — force download instead of inline view.
    """
    row = db.get_token_row(token)
    if row is None:
        abort(404)

    fmt = request.args.get("fmt", "pdf").lower()
    as_download = request.args.get("download") == "1"
    ref = row["proposal_ref"] or "proposal"

    if fmt == "docx":
        # The DOCX is saved alongside the PDF with the same file stem.
        docx_path = Path(row["pdf_path"]).with_suffix(".docx")
        if not docx_path.exists():
            abort(404)
        return send_file(
            str(docx_path),
            as_attachment=True,
            download_name=f"{ref}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    is_signed = bool(row["signed_pdf_path"]) and os.path.exists(row["signed_pdf_path"])
    pdf_path = row["signed_pdf_path"] if is_signed else row["pdf_path"]
    if not pdf_path or not os.path.exists(pdf_path):
        abort(404)
    return send_file(
        pdf_path,
        as_attachment=as_download,
        download_name=f"{ref}{'_signed' if is_signed else ''}.pdf",
        mimetype="application/pdf",
    )


@app.route("/api/sign/<token>", methods=["POST"])
def api_sign(token):
    """
    Accepts the client's signature and finalises the proposal.

    Body: {
      sig_type: "draw" | "type",
      sig_data: "<data:image/png;base64,...>" or "<typed name>",
      signer_name: str,
      signer_position: str (optional)
    }
    """
    row = db.get_token_row(token)
    valid, reason = db.is_token_valid(row)
    if not valid:
        return jsonify({"error": reason}), 400

    data = request.json or {}
    sig_type = data.get("sig_type")
    sig_data = data.get("sig_data")
    signer_name = (data.get("signer_name") or row["client_name"] or "").strip()
    signer_position = (data.get("signer_position") or "").strip()

    if sig_type not in ("draw", "type"):
        return jsonify({"error": "sig_type must be 'draw' or 'type'"}), 400
    if not sig_data:
        return jsonify({"error": "sig_data is required"}), 400
    if not signer_name:
        return jsonify({"error": "signer_name is required"}), 400

    signed_date = datetime.utcnow().strftime("%d %B %Y")
    signed_pdf_path = str(SIGNED_DIR / f"{row['proposal_ref']}_{token[:8]}_signed.pdf")

    try:
        signing.stamp_signature(
            row["pdf_path"],
            signed_pdf_path,
            sig_type=sig_type,
            sig_data=sig_data,
            signer_name=signer_name,
            signer_position=signer_position,
            signed_date=signed_date,
        )
    except Exception as e:
        print(f"[api_sign] stamping failed: {e}")
        return jsonify({"error": f"Could not apply signature: {e}"}), 500

    ip_address = request.headers.get("X-Forwarded-For", request.remote_addr)

    db.mark_signed(
        token,
        sig_type=sig_type,
        sig_data=sig_data if sig_type == "type" else "[drawn signature image]",
        signer_name=signer_name,
        signer_position=signer_position,
        signed_pdf_path=signed_pdf_path,
        ip_address=ip_address,
    )

    print(
        f"[api_sign] sig_type={sig_type} sig_data_len={len(sig_data) if sig_data else 0} signer_name={signer_name!r} signer_position={signer_position!r}"
    )

    # Notify the sender (firm) that the proposal has been signed
    sender_email = row.get("sender_email")
    if sender_email:
        signed_at_display = datetime.utcnow().strftime("%d %B %Y, %H:%M UTC")
        mail_ok, mail_err = mailer.send_signed_notification_email(
            sender_email=sender_email,
            client_name=row["client_name"],
            client_email=row["client_email"],
            proposal_ref=row["proposal_ref"],
            project_address=row["project_address"],
            signed_pdf_path=signed_pdf_path,
            signer_name=signer_name,
            signed_at=signed_at_display,
        )
        if not mail_ok:
            print(f"[api_sign] ⚠️ sender notification email failed: {mail_err}")
    else:
        print(
            "[api_sign] ⚠️ no sender_email stored for this token — skipping notification"
        )

    # Send the client their signed copy
    copy_ok, copy_err = mailer.send_client_signed_copy_email(
        client_name=row["client_name"],
        client_email=row["client_email"],
        proposal_ref=row["proposal_ref"],
        project_address=row["project_address"] or "",
        signed_pdf_path=signed_pdf_path,
        sender_name=row.get("sender_name") or "Concept Engineers Team",
        sender_title=row.get("sender_title") or "",
    )
    if not copy_ok:
        print(f"[api_sign] ⚠️ client signed-copy email failed: {copy_err}")

    return jsonify(
        {
            "success": True,
            "message": "Proposal signed successfully.",
            "downloadUrl": f"/document/{token}?download=1",
        }
    )


@app.route("/proposals")
@auth.login_required
def proposals_page():
    rows = db.list_recent(500)
    for r in rows:
        r["display_status"] = db.display_status(r)
        r["is_invite"] = False

    # Pending self-service quote invites appear alongside sent proposals.
    # (Completed invites already show as signed proposals, so only pending ones.)
    for inv in db.list_pending_invites(200):
        rows.append(
            {
                "token": inv["token"],
                "proposal_ref": "Quote invite",
                "client_name": inv["client_name"],
                "client_email": inv["client_email"],
                "project_address": inv["project_address"] or "—",
                "created_at": inv["created_at"],
                "signed_at": None,
                "display_status": db.invite_display_status(inv),
                "is_invite": True,
            }
        )

    rows.sort(key=lambda r: r.get("created_at") or "", reverse=True)

    # Compute the summary tiles from the SAME merged list the table renders,
    # so the counts always match what's on screen. (db.get_stats() only counted
    # signing_tokens and missed the pending quote invites, so TOTAL SENT read
    # lower than the number of visible rows.)
    stats = {"total": len(rows), "pending": 0, "signed": 0, "expired": 0, "voided": 0}
    for r in rows:
        st = r["display_status"]
        if st in ("pending", "invited"):
            stats["pending"] += 1  # awaiting client action (sign OR build+sign)
        elif st in ("signed", "quote_accepted"):
            stats["signed"] += 1  # completed / accepted
        elif st == "expired":
            stats["expired"] += 1
        elif st == "void":
            stats["voided"] += 1

    return render_template("proposals.html", proposals=rows, stats=stats)


@app.route("/api/resend/<token>", methods=["POST"])
@auth.login_required
def resend_proposal(token):
    row = db.get_token_row(token)
    if row is None:
        return jsonify({"error": "Proposal not found"}), 404
    if row["status"] == "signed":
        return jsonify({"error": "This proposal has already been signed"}), 400
    if row["status"] == "void":
        return jsonify({"error": "This signing link has been cancelled"}), 400

    signing_url = f"{PUBLIC_BASE_URL.rstrip('/')}/sign/{row['token']}"
    mail_ok, mail_err = mailer.send_signing_link_email(
        client_name=row["client_name"],
        client_email=row["client_email"],
        proposal_ref=row["proposal_ref"],
        project_address=row["project_address"],
        signing_url=signing_url,
        sender_name=row.get("sender_name") or "Concept Engineers Team",
        sender_title=row.get("sender_title") or "",
    )
    if mail_ok:
        return jsonify(
            {
                "success": True,
                "message": f"Signing link resent to {row['client_email']}",
            }
        )
    return jsonify({"error": mail_err}), 500


@app.route("/api/smart-quote-sign", methods=["POST"])
def smart_quote_sign():
    """
    Client-submitted smart quote with inline signature.
    Generates DOCX → PDF → stamps signature → saves to DB → emails CE team.
    """
    data = request.json or {}

    scope_id = (data.get("scope") or "").strip()
    size_id = (data.get("size") or "medium").strip()
    tier_id = (data.get("tier") or "mid").strip()
    extras = data.get("extras") or []
    fee_low = int(data.get("fee_low") or 0)
    fee_high = int(data.get("fee_high") or 0)
    fee_mid = int(data.get("fee_mid") or ((fee_low + fee_high) // 2))

    client_name = (data.get("client_name") or "").strip()
    client_email = (data.get("client_email") or "").strip().lower()
    company = (data.get("company") or "").strip()
    phone = (data.get("phone") or "").strip()
    project_address = (data.get("project_address") or "").strip()
    sig_type = (data.get("sig_type") or "").strip()
    sig_data_val = data.get("sig_data") or ""
    invite_token = (data.get("invite_token") or "").strip()

    # The builder is only reachable via an emailed invite link — validate it.
    invite_row = db.get_quote_invite(invite_token) if invite_token else None
    invite_ok, invite_reason = db.is_invite_valid(invite_row)
    if not invite_ok:
        return (
            jsonify({"error": invite_reason or "A valid quote invite is required"}),
            400,
        )

    if not client_name:
        return jsonify({"error": "client_name is required"}), 400
    if not client_email:
        return jsonify({"error": "client_email is required"}), 400
    if not project_address:
        return jsonify({"error": "project_address is required"}), 400
    if sig_type not in ("draw", "type"):
        return jsonify({"error": "Valid signature is required"}), 400
    if not sig_data_val:
        return jsonify({"error": "sig_data is required"}), 400

    ref = "CEQ-" + datetime.utcnow().strftime("%Y%m%d%H%M%S")
    today_str = datetime.utcnow().strftime("%d %B %Y")

    # ── Map quote selections to the template's scope switches ───────────────
    is_da = scope_id in ("da", "full")
    is_dd = scope_id in ("dd", "full")
    is_uu_minor = scope_id == "uu"
    is_unitywater = scope_id == "uw_only"
    has_construction = "construction" in extras
    has_as_constructed = "as_constructed" in extras

    # Human project type from the quote's own selector (fixes the earlier
    # bug where this client answer was collected but discarded).
    PROJECT_TYPE_LABELS = {
        "residential": "Residential Subdivision",
        "commercial": "Commercial",
        "industrial": "Industrial",
        "service_station": "Service Station",
    }
    project_type = PROJECT_TYPE_LABELS.get(
        (data.get("project_type") or "").strip(), "Residential Subdivision"
    )

    TIER_LABELS = {
        "bronze": "Bronze",
        "silver": "Silver",
        "gold": "Gold",
        "platinum": "Platinum",
    }
    tier_label = TIER_LABELS.get(tier_id, "Gold")

    # ── Live-computed fees from the quote's pricing engine ──────────────────
    # Phase add-on prices — keep in sync with EXTRAS in smart_quote.html.
    CONSTRUCTION_PRICE = 2200.0
    AS_CONSTRUCTED_PRICE = 950.0

    # The quote's fee_mid covers scope + size/tier adjustments + all extras.
    # Carve out the phase add-ons (they get their own fee-table rows); the
    # remaining amount is the primary scope fee, split DA:DD proportionally
    # to their base prices for the Full Civil Package.
    phases_amount = (CONSTRUCTION_PRICE if has_construction else 0) + (
        AS_CONSTRUCTED_PRICE if has_as_constructed else 0
    )
    scope_amount = max(float(fee_mid) - phases_amount, 0.0)

    fees = dict(REFERENCE_FEES)
    fees["dd_construction_fee"] = CONSTRUCTION_PRICE
    fees["dd_as_constructed_fee"] = AS_CONSTRUCTED_PRICE
    fees["uu_minor_construction_fee"] = CONSTRUCTION_PRICE
    fees["uu_minor_as_constructed_fee"] = AS_CONSTRUCTED_PRICE

    if scope_id == "full":
        da_share = 7200.0 / (7200.0 + 9500.0)
        fees["da_fee"] = round(scope_amount * da_share, 2)
        fees["dd_fee"] = round(scope_amount - fees["da_fee"], 2)
    elif scope_id == "da":
        fees["da_fee"] = scope_amount
    elif scope_id == "dd":
        fees["dd_fee"] = scope_amount
    elif scope_id == "uu":
        fees["uu_minor_fee"] = scope_amount
    elif scope_id == "uw_only":
        fees["unitywater_fee"] = scope_amount

    ctx = build_template_context(
        contact_name=client_name,
        company_name=company,
        project_address=project_address,
        contact_email=client_email,
        contact_phone=phone,
        project_type=project_type,
        proposal_ref=ref,
        revision="Rev A",
        proposal_date=today_str,
        tier_label=tier_label,
        is_da=is_da,
        is_dd=is_dd,
        is_uu_minor=is_uu_minor,
        is_uu_major=False,
        is_unitywater=is_unitywater,
        has_construction=has_construction,
        has_as_constructed=has_as_constructed,
        fees=fees,
        sender_name="Concept Engineers Team",
        sender_title="",
        sender_email="admin@conceptengineers.com.au ",
        sender_phone="+61 7 3505 6498",
    )

    tpl = DocxTemplate(str(TEMPLATE))
    tpl.render(ctx)

    safe_ref = ref.replace("-", "_")
    tmp_docx_path = PROPOSALS_DIR / f"{safe_ref}.docx"
    pdf_path = PROPOSALS_DIR / f"{safe_ref}.pdf"
    signed_pdf_path = SIGNED_DIR / f"{safe_ref}_signed.pdf"

    tpl.save(str(tmp_docx_path))
    inject_deferred_signing_tags(str(tmp_docx_path))

    converted, convert_err = docx_to_pdf(str(tmp_docx_path), str(pdf_path))
    if not converted:
        return jsonify({"error": f"Could not generate PDF: {convert_err}"}), 500

    # Stamp signature — if this fails, the client must NOT be told the
    # proposal was signed successfully (that would be a false positive: the
    # DB would show "signed" and the client would receive an unsigned PDF
    # captioned as their signed copy).
    signed_date = datetime.utcnow().strftime("%d %B %Y")
    try:
        signing.stamp_signature(
            str(pdf_path),
            str(signed_pdf_path),
            sig_type=sig_type,
            sig_data=sig_data_val,
            signer_name=client_name,
            signer_position=company,
            signed_date=signed_date,
        )
    except Exception as e:
        print(f"[smart_quote_sign] Signature stamp failed: {e}")
        return jsonify({"error": f"Could not apply signature: {e}"}), 500

    # Save to DB: create token then immediately mark signed
    ip_address = request.headers.get("X-Forwarded-For", request.remote_addr)
    token_info = db.create_signing_token(
        proposal_ref=ref,
        client_name=client_name,
        client_email=client_email,
        project_address=project_address,
        pdf_path=str(pdf_path),
        sender_name="Concept Engineers Team",
        sender_title="",
        sender_email="admin@conceptengineers.com.au",
        source="quote",
    )
    db.mark_signed(
        token_info["token"],
        sig_type=sig_type,
        sig_data=sig_data_val if sig_type == "type" else "[drawn signature image]",
        signer_name=client_name,
        signer_position=company,
        signed_pdf_path=str(signed_pdf_path),
        ip_address=ip_address,
    )

    # Email CE team
    mail_ok, mail_err = mailer.send_signed_notification_email(
        sender_email="admin@conceptengineers.com.au ",
        client_name=client_name,
        client_email=client_email,
        proposal_ref=ref,
        project_address=project_address,
        signed_pdf_path=str(signed_pdf_path),
        signer_name=client_name,
        signed_at=datetime.utcnow().strftime("%d %B %Y, %H:%M UTC"),
    )
    if not mail_ok:
        print(f"[smart_quote_sign] ⚠️ team notification email failed: {mail_err}")

    # Send the client their signed copy
    copy_ok, copy_err = mailer.send_client_signed_copy_email(
        client_name=client_name,
        client_email=client_email,
        proposal_ref=ref,
        project_address=project_address,
        signed_pdf_path=str(signed_pdf_path),
    )
    if not copy_ok:
        print(f"[smart_quote_sign] ⚠️ client signed-copy email failed: {copy_err}")

    # Close out the invite so the link can't be reused
    db.mark_invite_completed(invite_token, ref)

    print(f"[smart_quote_sign] ✅ ref={ref} client={client_email}")

    return jsonify(
        {
            "success": True,
            "ref": ref,
            "downloadUrl": f"/document/{token_info['token']}?download=1",
            "emailedCopy": copy_ok,
            "message": f"Proposal signed. Ref: {ref}",
        }
    )


if __name__ == "__main__":
    app.run(debug=True, port=5000)
