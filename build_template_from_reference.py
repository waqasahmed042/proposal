"""
Transforms master_proposal_template.docx (a fresh copy of the client's real
26000-FP01-Proposal.docx) into a docxtpl-compatible Jinja template.

Key docxtpl mechanics this relies on (verified against docxtpl 0.20.2 source):
- A `{%tr ... %}` tag consumes its ENTIRE table row and is replaced by the
  bare Jinja tag — so each tr tag must live in its own dedicated row.
- A `{%p ... %}` tag likewise consumes its whole paragraph — dedicated
  control paragraphs self-remove and leave no blank lines behind.
"""
import copy
from docx import Document
from docx.oxml.ns import qn

PATH = r"e:\Projects\Concept Engineers\fee-proposal-automation\master_proposal_template.docx"

doc = Document(PATH)
body = doc.element.body


# ---------- helpers ----------

def set_text_single_run(paragraph, new_text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def set_text_two_runs(paragraph, label_text, value_text):
    """Keep run 0 as the (bold) label, put value in the last run, blank the middle."""
    runs = paragraph.runs
    if len(runs) >= 2:
        runs[0].text = label_text
        runs[-1].text = value_text
        for r in runs[1:-1]:
            r.text = ""
    else:
        set_text_single_run(paragraph, label_text + value_text)


def new_control_paragraph(text):
    p = doc.add_paragraph(text)
    body.remove(p._p)
    return p._p


def insert_before(target_element, text):
    target_element.addprevious(new_control_paragraph(text))


def insert_after(target_element, text):
    target_element.addnext(new_control_paragraph(text))


def set_cell_text(cell, text):
    set_text_single_run(cell.paragraphs[0], text)


def set_cell_fee_dual(cell, text):
    """Fee cells stack a strikethrough 'was' price over the real discounted
    price. Write into the first NON-strikethrough paragraph so the rendered
    fee doesn't appear crossed out; blank the rest."""
    paras = cell.paragraphs
    target = None
    for p in paras:
        runs = p.runs
        struck = bool(runs) and runs[0].font.strike
        if not struck:
            target = p
            break
    if target is None:
        target = paras[-1]
    for p in paras:
        set_text_single_run(p, text if p is target else "")


def set_cell_total(cell, text):
    """Total cells: ['$X.00 ', '(Ex GST)'] — replace price, keep '(Ex GST)'."""
    set_text_single_run(cell.paragraphs[0], text)


def make_tag_row(template_row, tag_text):
    """Deepcopy a row, blank all text, put tag_text in the first w:t.
    docxtpl consumes the whole row, so formatting of the copy is irrelevant."""
    new_tr = copy.deepcopy(template_row._tr)
    ts = list(new_tr.iter(qn("w:t")))
    for t in ts:
        t.text = ""
    if ts:
        ts[0].text = tag_text
    return new_tr


def wrap_row_conditional(row, condition_expr):
    """Give `row` a {%tr if %} row above and {%tr endif %} row below."""
    row._tr.addprevious(make_tag_row(row, f"{{%tr if {condition_expr} %}}"))
    row._tr.addnext(make_tag_row(row, "{%tr endif %}"))


P = doc.paragraphs

# ============================================================
# 0. DATA-BOUND CONTENT CONTROLS (w:sdt)
#
# The client's document binds cover/letter fields to customXml/item8.xml
# (FeeProposalInputs: Address / ClientName / ProjectType). Word REFRESHES
# bound controls from that XML part on open — which would overwrite our
# rendered values during docx2pdf conversion. So: put the Jinja tag into
# each control's cached text and STRIP the binding.
#
# Controls found (host): TABLE0 K-table values ×3, P[63] cover address,
# P[75] cover client name, P[94] intro sentence ×2 (inline!), P[393] team.
# ============================================================
SDT_MAP = {
    "Address": "{{ project_address }}",
    "ClientName": "{{ contact_name }}",
    "ProjectType": "{{ project_type }}",
}
for sdt in list(body.iter(qn("w:sdt"))):
    pr = sdt.find(qn("w:sdtPr"))
    if pr is None:
        continue
    binding = pr.find(qn("w:dataBinding"))
    if binding is None:
        continue
    xpath = binding.get(qn("w:xpath")) or ""
    tag = next((v for k, v in SDT_MAP.items() if k in xpath), None)
    if tag is None:
        continue
    ts = list(sdt.iter(qn("w:t")))
    for t in ts:
        t.text = ""
    if ts:
        ts[0].text = tag
    pr.remove(binding)

# ============================================================
# 1. COVER PAGE
#
# P[74] "Prepared for :" label and P[75] (ClientName sdt) stay as designed.
# Ref goes into the spare empty line P[76]; date P[77] becomes dynamic.
# ============================================================
set_text_single_run(P[76], "Ref: {{ proposal_ref }}  |  {{ revision }}")
set_text_single_run(P[77], "{{ proposal_date }}")

# ============================================================
# 2. COVER INFO TABLE (table 0)
# Rows 0-2 (Address / Client Name / Project Type) already carry bound sdt
# value controls — handled by step 0. Only the empty offering cells here.
# ============================================================
t0 = doc.tables[0]
set_cell_text(t0.rows[3].cells[1], "{{ tier_label }}")
set_cell_text(t0.rows[4].cells[1], "{{ da_dd_label }}")
set_cell_text(t0.rows[5].cells[1], "{{ uu_uw_label }}")

# ============================================================
# 3. INTRO LETTER
# P[94]'s sentence contains INLINE bound sdts for project type + address —
# already converted by step 0, so the sentence needs no rewrite.
# ============================================================
set_text_single_run(P[102], "{{ sender_name }}\t\t")
set_text_single_run(P[103], "{{ sender_title }}")
set_text_two_runs(P[105], "E \t", "{{ sender_email }}")
set_text_two_runs(P[106], "M\t", "{{ sender_phone }}")
set_text_two_runs(P[107], "L\t", "(07) 3505 6498")

# ============================================================
# 4. TIER MENTION (paragraph 460)
# ============================================================
P[460].runs[10].text = "{{ tier_label }}"
P[460].runs[11].text = ""

# ============================================================
# 5. SCOPE CONDITIONALS — DA / DD timeline + value blocks
# NOTE: all inserts done bottom-up relative to paragraph objects captured
# BEFORE any insertion, so python-docx element refs stay valid.
# ============================================================
insert_before(P[419]._p, "{%p if is_da %}")
insert_before(P[439]._p, "{%p endif %}")

insert_before(P[439]._p, "{%p if is_dd %}")
insert_before(P[458]._p, "{%p endif %}")

insert_before(P[458]._p, "{%p if is_da or is_dd %}")
insert_before(P[545]._p, "{%p endif %}")

# ============================================================
# 6. FEE TABLES
#
# Body order in the fee region (one Word section, P544 sectPr .. P597 sectPr):
#   P545-549 DA heading+intro | TABLE1 | P550-554 DD heading+intro | TABLE2 |
#   P555-556 | P557-558 UUmin intro | TABLE3 | P559-576 spacers |
#   P577-578 UUmaj intro | TABLE4 | P579-594 spacers | P595-596 UW intro |
#   TABLE5 | P597 sectPr
#
# The original relies on runs of empty spacer paragraphs to push each fee
# table onto its own page. Removed blocks would leave those spacers behind
# as header-only blank pages, so instead: delete the spacer runs and give
# each block's first paragraph an explicit page-break-before.
# ============================================================
t1, t2, t3, t4, t5 = doc.tables[1:6]

# Delete the big spacer runs (fragile page filler)
for idx in list(range(559, 577)) + list(range(579, 595)):
    P[idx]._p.getparent().remove(P[idx]._p)

# --- DA (table 1) ---
insert_before(P[545]._p, "{%p if is_da %}")
set_cell_fee_dual(t1.rows[1].cells[1], "{{ da_fee }}")
set_cell_text(t1.rows[2].cells[1], "Hourly Rates")
insert_after(t1._tbl, "{%p endif %}")

# --- Detailed Design (table 2) ---
# Capture row refs BEFORE structural changes (tag-row inserts shift indices)
t2_base, t2_rfi, t2_constr = t2.rows[1], t2.rows[2], t2.rows[3]
t2_alt1, t2_alt2, t2_alt3 = t2.rows[4], t2.rows[5], t2.rows[6]
t2_ascon, t2_post, t2_total = t2.rows[7], t2.rows[8], t2.rows[9]

insert_before(P[550]._p, "{%p if is_dd %}")
P[550].paragraph_format.page_break_before = True
set_cell_fee_dual(t2_base.cells[1], "{{ dd_fee }}")
set_cell_text(t2_rfi.cells[1], "Hourly Rates")
set_cell_text(t2_constr.cells[1], "{{ dd_construction_fee }}")
set_cell_text(t2_ascon.cells[1], "{{ dd_as_constructed_fee }}")
set_cell_text(t2_post.cells[1], "Hourly Rates")
set_cell_total(t2_total.cells[1], "{{ dd_total }}")
# Remove the D&C alternative construction rows (kept: the standard one)
for r in (t2_alt1, t2_alt2, t2_alt3):
    r._tr.getparent().remove(r._tr)
wrap_row_conditional(t2_constr, "has_construction")
wrap_row_conditional(t2_ascon, "has_as_constructed")
insert_after(t2._tbl, "{%p endif %}")

# --- UU Minor Works (table 3) ---
t3_base, t3_constr, t3_supp = t3.rows[1], t3.rows[2], t3.rows[3]
t3_ascon, t3_post, t3_total = t3.rows[4], t3.rows[5], t3.rows[6]

insert_before(P[557]._p, "{%p if is_uu_minor %}")
P[557].paragraph_format.page_break_before = True
set_cell_text(t3_base.cells[1], "{{ uu_minor_fee }}")
set_cell_text(t3_constr.cells[1], "{{ uu_minor_construction_fee }}")
set_cell_text(t3_supp.cells[1], "Hourly Rates")
set_cell_text(t3_ascon.cells[1], "{{ uu_minor_as_constructed_fee }}")
set_cell_text(t3_post.cells[1], "Hourly Rates")
set_cell_total(t3_total.cells[1], "{{ uu_minor_total }}")
wrap_row_conditional(t3_constr, "has_construction")
wrap_row_conditional(t3_ascon, "has_as_constructed")
insert_after(t3._tbl, "{%p endif %}")

# --- UU Major Works (table 4) ---
t4_base, t4_supp, t4_ascon = t4.rows[1], t4.rows[2], t4.rows[3]
t4_post, t4_total = t4.rows[4], t4.rows[5]

insert_before(P[577]._p, "{%p if is_uu_major %}")
P[577].paragraph_format.page_break_before = True
set_cell_text(t4_base.cells[1], "{{ uu_major_fee }}")
set_cell_text(t4_supp.cells[1], "Hourly Rates")
set_cell_text(t4_ascon.cells[1], "{{ uu_major_as_constructed_fee }}")
set_cell_text(t4_post.cells[1], "Hourly Rates")
set_cell_total(t4_total.cells[1], "{{ uu_major_total }}")
wrap_row_conditional(t4_ascon, "has_as_constructed")
insert_after(t4._tbl, "{%p endif %}")

# --- UnityWater (table 5) ---
insert_before(P[595]._p, "{%p if is_unitywater %}")
P[595].paragraph_format.page_break_before = True
set_cell_text(t5.rows[1].cells[1], "{{ unitywater_fee }}")
set_cell_text(t5.rows[2].cells[1], "Hourly Rates")
set_cell_total(t5.rows[3].cells[1], "{{ unitywater_total }}")
insert_after(t5._tbl, "{%p endif %}")

# ============================================================
# 6b. TERMS PAGE SIGN-OFF (right column, above the Authorisation table)
# Hardcoded "Zac Lemon / Director..." becomes the dynamic sender. The two
# blank spacer lines are removed and the title line is conditional — the
# filled Authorisation cells wrap to 2 lines each (company, address), which
# otherwise pushes the Print Name / Date rows onto the back-cover photo.
# ============================================================
set_text_single_run(P[626], "{{ sender_name }}")
set_text_single_run(P[627], "{{ sender_title }}")
insert_before(P[627]._p, "{%p if sender_title %}")
insert_after(P[627]._p, "{%p endif %}")
# P620/P623 are additional single blank lines in the same column — removed
# so the Date row (which carries the [[date]] signing placeholder) cannot
# land on the back-cover photo page where the stamp would be invisible.
for spacer in (P[620], P[623], P[625], P[628]):
    spacer._p.getparent().remove(spacer._p)

# Same conditional treatment for the letter-page title line (P[103]) so an
# empty sender title doesn't leave a stray blank line there either.
insert_before(P[103]._p, "{%p if sender_title %}")
insert_after(P[103]._p, "{%p endif %}")

# ============================================================
# 7. AUTHORISATION TO PROCEED (table 7)
# ============================================================
t7 = doc.tables[7]
set_cell_text(t7.rows[1].cells[1], "{{ proposal_ref }}")
set_cell_text(t7.rows[2].cells[1], "{{ company_name }}")
set_cell_text(t7.rows[3].cells[1], "{{ project_address }}")
set_cell_text(t7.rows[4].cells[1], "{{ contact_phone }}")
set_cell_text(t7.rows[5].cells[1], "{{ contact_email }}")
set_cell_text(t7.rows[6].cells[1], "{{ accounts_phone }}")
set_cell_text(t7.rows[7].cells[1], "{{ accounts_email }}")
set_cell_text(t7.rows[8].cells[1], "[[signature]]")
set_cell_text(t7.rows[9].cells[1], "{{ contact_name }}")
set_cell_text(t7.rows[10].cells[1], "[[date]]")

doc.save(PATH)
print("Template transformation complete.")
