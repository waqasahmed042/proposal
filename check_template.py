"""
check_template.py — run this once to see exactly what's in
master_proposal_template.docx around the signature/date placeholders.

Usage (from your project root, same folder as app.py):
    python check_template.py
"""

from docx import Document

TEMPLATE_PATH = "master_proposal_template.docx"


def scan_paragraphs(paragraphs, location_label):
    found_any = False
    for i, p in enumerate(paragraphs):
        text = p.text
        if (
            "signature" in text.lower()
            or "{{date" in text.lower()
            or "[[date" in text.lower()
            or "date}}" in text.lower()
            or "date]]" in text.lower()
        ):
            found_any = True
            print(f"[{location_label} paragraph {i}] raw text:")
            print(f"    {text!r}")
            print(f"    (runs: {[r.text for r in p.runs]})")
            print()
    return found_any


def main():
    doc = Document(TEMPLATE_PATH)

    print(f"Scanning {TEMPLATE_PATH} ...\n")

    found = scan_paragraphs(doc.paragraphs, "body")

    for t_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                if scan_paragraphs(cell.paragraphs, f"table {t_idx} cell"):
                    found = True

    if not found:
        print("No paragraph containing 'signature' or 'date' was found at all.")
        print("Either the wording is different than expected, or these lines")
        print("are inside a text box / header / footer, which python-docx")
        print("does not scan by default.")
        return

    print("---")
    print("What to look for above:")
    print("  - If you see '{{signature}}' or '{{date}}' (curly braces): the")
    print("    template has NOT been updated yet. Change these to")
    print("    '[[signature]]' and '[[date]]' (square brackets) and re-save.")
    print("  - If you see '[[signature]]' / '[[date]]' but split oddly across")
    print("    multiple runs with extra characters in between (e.g. autocorrect")
    print("    inserted something), that's also worth fixing directly in Word.")
    print("  - If you see exactly '[[signature]]' and '[[date]]' as contiguous")
    print("    text, the template is correct and the issue is elsewhere (e.g.")
    print("    Flask still running old cached code — restart the server).")


if __name__ == "__main__":
    main()
